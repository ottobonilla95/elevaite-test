import docx

class DocxReader:
    def __init__(self, doc_path):
        self.doc = docx.Document(doc_path)
        pass

    def read(self):
        doc = self.doc
        text = ""
        for para in doc.paragraphs:
            text += "\n" + para.text
        return text

    # Function to check if a paragraph is a heading
    def is_heading(self, paragraph):
        # Check if the style name contains "Heading"
        return paragraph.style.name.startswith('Heading 1')

    def is_text(self, paragraph):
        return paragraph.style.name.startswith('Normal')


    # Function to extract headings and text separately
    def extract_headings_and_bodies(self,limit=None):
        doc = self.doc

        headings = []
        body_text = []

        # Loop through each paragraph in the document
        last_paragraph = ""
        heading_detected = False
        article_text = ""
        for para in doc.paragraphs[:limit]:
            if self.is_heading(para):
                heading_detected = True
            if not heading_detected:
                continue

            if self.is_heading(para):
                headings.append(para.text)
                if article_text!="":
                    article_text = article_text.replace(last_paragraph, "")
                    last_paragraph = ""
                    body_text.append(article_text)
                article_text = ""
            elif heading_detected:
                article_text += "\n"+para.text
                last_paragraph = para.text
        body_text.append(article_text)
        return headings, body_text


